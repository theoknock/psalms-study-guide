import openai
import os
import sys
import docx
from docx import Document

psalms = ("40", "10", "12", "23", "35", "38", "41", "88", "139", "141")

def style_subtitle(doc, response):
    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.style = doc.styles['SubtitlePS']
    subtitle_paragraph.add_run(str(response))

def prompts(psalm):
  # print(my_array[0]['function'](response))
  # my_array = [
  #   {'function': 'style_subtitle', 'prompt': f"Write a 5 to 6 sentence introduction to the spiritual and emotional elements of Psalm {psalm}, and try to connect its meaning and purpose with that of Christian faith."},
  #   {'function': 'NumberedPS',     'prompt': f"Provide the full text of Psalm {psalm} using the King James Version, formatted for organizational clarity."},
  #   {'function': 'BodyPS',         'prompt': f"Create a study guide for Psalm {psalm} that includes both a brief, one-sentence summary of each main point along with the Bible verses they correspond to (be sure to cover every verse in the psalm; do not call the summary “main point”; include the verses being summarized at the end), an extended summary for each main point (do not label the extended summary), two or three questions that challenge the reader's understanding of each main point — three is preferred (label the challenge questions, “EVALUATE”; omit the colon), and two or three questions that ask the reader to reflect on their own life experience for each main point  — three is preferred (label the reflection questions, “REFLECT”; omit the colon). Provide answers to the challenge questions, but not the questions for reflection."},
  #   {'function': 'BodyPS',         'prompt': f"Describe all the ways Psalm {psalm} embodies or reflects God’s nature in two sections: 1. divine (incommunicable) attributes; and, 2. communicable  attributes. Include biblical references, if applicable."}
  # ]
  return (
    f"Write a 5 to 6 sentence introduction to the spiritual and emotional elements of Psalm {psalm}, and try to connect its meaning and purpose with that of Christian faith.",
    f"Provide the full text of Psalm {psalm} using the King James Version, formatted for organizational clarity.",
    f"Create a study guide for Psalm {psalm} that includes both a brief, one-sentence summary of each main point along with the Bible verses they correspond to (be sure to cover every verse in the psalm; do not call the summary “main point”; include the verses being summarized at the end), an extended summary for each main point (do not label the extended summary), two or three questions that challenge the reader's understanding of each main point — three is preferred (label the challenge questions, “EVALUATE”; omit the colon), and two or three questions that ask the reader to reflect on their own life experience for each main point  — three is preferred (label the reflection questions, “REFLECT”; omit the colon). Provide answers to the challenge questions, but not the questions for reflection.",
    f"Describe all the ways Psalm {psalm} embodies or reflects God’s nature in two sections: 1. divine (incommunicable) attributes; and, 2. communicable  attributes. Include biblical references, if applicable.",
    f"Relate the person and teachings of Jesus Christ and Psalm {psalm}, particularly, as the pertain to the gospel. Include supporting Bible verses for every connection made, especially if there is a match between the words of Jesus and verses in this psalm.",
    f"List all of the psalms that are identical or highly similar to Psalm {psalm}, whether in part or in whole. Explain the similarities in as much detail as possible."
  )



try:
  openai.api_key = os.environ['OPENAI_API_KEY']
except KeyError:
  sys.stderr.write("""
  You haven't set up your API key: OPENAI_API_KEY.
  """)
  exit(1)

for psalm in psalms:
  filename = str(f"/Users/xcodedeveloper/Desktop/pip_install/ChatGPT Responses/Psalm_{psalm}.docx")
  if os.path.isfile(filename):
    print(f"File Psalm_{psalm}.docx exists — skipping...\n")
  else:
    doc = docx.Document('/Users/xcodedeveloper/Desktop/pip_install/Psalms_Template.docx')
    title_paragraph = doc.add_paragraph()
    title_paragraph.style = doc.styles['TitlePS']
    title_paragraph.add_run(str(psalm))
                            
    print(f"Exporting ChatGPT responses to Psalm_{psalm}.docx...\n")

    for prompt in prompts(psalm):
      response = openai.ChatCompletion.create(model="gpt-4",
                                              messages=[{
                                                "role": "user",
                                                "content": prompt
                                              }],
                                              temperature=1,
                                              top_p=1,
                                              n=3)
      print(str(f"Prompt: {prompt}") + '\n\n')
      print(str(f"Response: {response.choices[0].message['content']}") + '\n\n-------------------------\n\n')
      
      prompt_paragraph       = doc.add_paragraph(str(prompt))
      response_paragraph     = doc.add_paragraph(str(response.choices[0].message.content))
      prompt_paragraph.style = doc.styles['Heading3PS']
      response_paragraph     = doc.styles['BodyPS']
      # prompt_paragraph.add_run(str(prompt))
      # response_paragraph.add_run(str(response.choices[0].message.content))

      # with open(f"/Users/xcodedeveloper/Desktop/pip_install/ChatGPT Responses/{filename}", "a") as file:
      #   file.write(str(response.choices[0].message.content) + '\n\n-------------------------\n\n')
        
    doc.save(filename)
